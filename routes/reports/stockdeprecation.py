from fastapi import FastAPI, BackgroundTasks, UploadFile, File, Form, APIRouter, status, HTTPException, Depends
from starlette.responses import JSONResponse
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig, MessageType
from pydantic import EmailStr, BaseModel
from typing import List
from sqlalchemy.orm import Session
from database import models
from utils import schemas, utils, oauth2
from database.database import get_db
from docx import Document

class EmailSchema(BaseModel):
    email: List[EmailStr]


conf = ConnectionConfig(
    MAIL_USERNAME = "AKIAY74ITD5OE3BRFNE3",
    MAIL_PASSWORD = "BKEkY0Gr0522LV7vNWhcmk0cdrRlknCqpiSZ/2lOTBaG",
    MAIL_FROM = "niyobern@icloud.com",
    MAIL_PORT = 2587,
    MAIL_SERVER = "email-smtp.us-east-2.amazonaws.com",
    MAIL_FROM_NAME="Niyomugabo Bernard",
    MAIL_STARTTLS = True,
    MAIL_SSL_TLS = False,
    USE_CREDENTIALS = True,
    VALIDATE_CERTS = True
)

router = APIRouter(prefix="/stockdeprecation", tags=["Reports"])

async def make_document(title, data, email):
    a = data[0]
    titles = [x for x in a.keys()]
    colums = len(titles)
    document = Document()
    document.add_heading(title, 1)
    table = document.add_table(rows=1, cols=colums)
    header_cells = table.rows[0].cells
    for i in range(len(header_cells)):
        header_cells[i].text = titles[i]
    for item in data:
        row_cells = table.add_row().cells
        values = [x for x in item.values()]
        for index in range(len(titles)):
            row_cells[index].text = str(values[index])
    table.style = 'Colorful Grid Accent 1'
    document.save("Report.docx")
    filei = open("Report.docx", "rb")
    file = UploadFile(filename="report.docx", file=filei)
    message = MessageSchema(
            subject=title,
            recipients=[email],
            body="Your Report",
            subtype=MessageType.html,
            attachments=[{"file": file}])

    fm = FastMail(conf)

    await fm.send_message(message)
    filei.close()
    return "done"

@router.post("/")
async def send_file(
    background_tasks: BackgroundTasks,
    email:EmailStr,db: Session = Depends(get_db), current_user: schemas.User = Depends(
    oauth2.get_current_user), limit: int = 100, skip: int = 100, start: str = "2022-12-18", end: str = "2022-12-30"
    ) -> JSONResponse:
    if current_user.role.value == "no_role":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed")
    if current_user.role.value not in ("manager", "boss", "deputy_boss", "store_keeper"):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Not enough priviledge")
    items = db.query(models.StockDeprecation).filter(models.StockDeprecation.created_at.between(start, end)).limit(limit).offset(skip).all()
    items_info = []
    for item in items:
        stock_item = db.query(models.StockItem).filter(models.StockItem.id == item.stock_id).first()
        created_time = str(item.created_at)
        info = {"id": item.id, "stock_id": stock_item.name, "quantity": item.quantity, "created_at": created_time[:10], "tag": item.tag}
        items_info.append(info)
    background_tasks.add_task(make_document,"Removal of Items from Stock",items_info,email)

    return JSONResponse(status_code=200, content={"message": "Your report is being processed"})

