from docx import Document
import io
from docx.shared import Inches
def make_document(title, data):
    file = io.BytesIO()
    rows = len(data)
    a = data[0]
    titles = [x for x in a.keys()]
    colums = len(titles)
    document = Document()
    document.add_heading(title, 1)
    table = document.add_table(rows=1, cols=colums)
    header_cells = table.rows[0].cells
    for i in range(len(header_cells)):
        header_cells[i].text = titles[i]
    for item in data:
        row_cells = table.add_row().cells
        values = [x for x in item.values()]
        for index in range(len(titles)):
            row_cells[index].text = str(values[index])
    table.style = 'Colorful Grid Accent 1'
    document.save(file)
    return file

title = "Bank Statement"
data = [{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123},{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123},{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}]

file = make_document(title, data)

from docx import Document
new_file = Document(file)
new_file.save("bahoneza.docx")