from fastapi import FastAPI, BackgroundTasks, UploadFile, File, Form, APIRouter
from starlette.responses import JSONResponse
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig, MessageType
from pydantic import EmailStr, BaseModel
from typing import List
from docx import Document

class EmailSchema(BaseModel):
    email: List[EmailStr]


conf = ConnectionConfig(
    MAIL_USERNAME = "AKIAY74ITD5OE3BRFNE3",
    MAIL_PASSWORD = "BKEkY0Gr0522LV7vNWhcmk0cdrRlknCqpiSZ/2lOTBaG",
    MAIL_FROM = "niyobern@icloud.com",
    MAIL_PORT = 2587,
    MAIL_SERVER = "email-smtp.us-east-2.amazonaws.com",
    MAIL_FROM_NAME="Niyomugabo Bernard",
    MAIL_STARTTLS = True,
    MAIL_SSL_TLS = False,
    USE_CREDENTIALS = True,
    VALIDATE_CERTS = True
)

router = APIRouter()

async def make_document(title, data, email):
    rows = len(data)
    a = data[0]
    titles = [x for x in a.keys()]
    colums = len(titles)
    document = Document()
    document.add_heading(title, 1)
    table = document.add_table(rows=1, cols=colums)
    header_cells = table.rows[0].cells
    for i in range(len(header_cells)):
        header_cells[i].text = titles[i]
    for item in data:
        row_cells = table.add_row().cells
        values = [x for x in item.values()]
        for index in range(len(titles)):
            row_cells[index].text = str(values[index])
    table.style = 'Colorful Grid Accent 1'
    document.save("Report.docx")
    filei = open("Report.docx", "rb")
    file = UploadFile(filename="report.docx", file=filei)
    message = MessageSchema(
            subject=title,
            recipients=[email],
            body="Your Report",
            subtype=MessageType.html,
            attachments=[{"file": file}])

    fm = FastMail(conf)

    await fm.send_message(message)
    filei.close()
    return "done"

@router.post("/")
async def send_file(
    background_tasks: BackgroundTasks,
    email:EmailStr = Form(...),
    ) -> JSONResponse:
    title = "Bank Statement"
    data = [{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123},{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123},{"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}, {"january": 200, "February": 405, "March": 2436, "June": 234, "September": 123}]
    
    background_tasks.add_task(make_document,title,data,email)


    return JSONResponse(status_code=200, content={"message": "Your report is being processed"})

